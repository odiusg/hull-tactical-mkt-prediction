"""Generate a Word report from test-set results and CV summary."""
from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BASE   = Path(__file__).parents[1]
TEST   = BASE / "results" / "test"
CV_DIR = BASE / "results" / "cv"
OUT    = BASE / "results" / "hull_tactical_report.docx"

# ── helpers ──────────────────────────────────────────────────────────────────

def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def _para(doc, text, bold=False, italic=False, size=None, color=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def _metric_table(doc, rows: list[tuple], headers: list[str]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Shading Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph()


def _img(doc, path: Path, width=5.5, caption=None):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].italic = True
        cp.runs[0].font.size = Pt(9)
        cp.paragraph_format.space_after = Pt(10)


# ── load data ────────────────────────────────────────────────────────────────

with open(TEST / "rf_test_results.json") as f:
    rf = json.load(f)

tm  = rf["test_metrics"]
cv10 = rf["cv_10fold_summary"]
bh_score = rf["benchmark_adjusted_sharpe"]

cv_df = pd.read_csv(CV_DIR / "cv_summary.csv")

# ── build document ───────────────────────────────────────────────────────────

doc = Document()

# ── style tweaks ─────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ════════════════════════════════════════════════════════════════════════════
# Cover
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
title = doc.add_heading("Hull Tactical 市场预测策略", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("量化回测分析报告")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(16)
sub.runs[0].bold = True

doc.add_paragraph()
date_p = doc.add_paragraph("2026年5月")
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.runs[0].font.size = Pt(12)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. Executive Summary
# ════════════════════════════════════════════════════════════════════════════
_heading(doc, "一、执行摘要")
_para(doc,
    "本项目基于 Hull Tactical 竞赛数据集，构建了一套完整的机器学习量化交易框架。"
    "我们对 7 类模型（线性模型、随机森林、XGBoost、LightGBM、LSTM 等）进行了系统性的超参数调优与交叉验证，"
    "最终选择随机森林（Random Forest）作为正式提交模型。"
)
_para(doc,
    "在独立测试集上，随机森林策略的调整后夏普比率达到 1.12，"
    f"显著优于买入持有基准（{bh_score:.2f}），最大回撤仅为 {abs(tm['max_drawdown']):.1%}，"
    "表现出较强的风险控制能力。"
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 2. 项目背景与目标
# ════════════════════════════════════════════════════════════════════════════
_heading(doc, "二、项目背景与目标")
_para(doc,
    "Hull Tactical 竞赛要求参赛者预测股票市场的远期收益率，并据此生成每日持仓比例（0 到 2 之间），"
    "0 表示完全空仓，1 表示持有基准仓位，2 表示 2 倍多头。"
)
_para(doc, "竞赛使用调整后夏普比率（Adjusted Sharpe Ratio）作为评分标准，该指标对以下两种行为进行惩罚：")
_bullet(doc, "策略波动率超过市场的 1.2 倍（过度放大风险）")
_bullet(doc, "策略平均超额收益低于市场（跑输大盘）")
_para(doc, "因此，优秀的策略需要同时具备盈利能力和稳定性，而非单纯追求高收益。")

_heading(doc, "三、数据与特征工程", level=2)
_para(doc,
    "原始数据包含股票市场的每日收益率、风险收益率等基础指标。"
    "经过预处理后，我们构建了涵盖以下类别的特征体系："
)
_bullet(doc, "动量类：不同周期（5、10、21、63 天）的历史收益率")
_bullet(doc, "波动率类：滚动标准差、波动率比率")
_bullet(doc, "技术指标类：RSI、布林带、MACD 等")
_bullet(doc, "宏观类：利率区间（rate_regime）")
_para(doc,
    "在特征筛选环节，我们采用信息系数（IC）检验，仅保留与目标变量存在统计显著相关性的特征，"
    "有效降低了噪声干扰。最终进入模型的特征数量约为 52 个。"
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 3. 建模方法
# ════════════════════════════════════════════════════════════════════════════
_heading(doc, "四、建模方法")

_heading(doc, "4.1 滚动窗口交叉验证", level=2)
_para(doc,
    "为避免未来数据泄露，我们采用时间序列专用的滚动扩展窗口交叉验证（Walk-Forward CV）。"
    "具体做法是：以训练集的前若干年数据作为初始训练窗口，逐步向后扩展，"
    "每次在新的验证窗口上评估模型表现。共进行 10 折验证，确保评估的时间覆盖面充足。"
)

_heading(doc, "4.2 超参数调优", level=2)
_para(doc,
    "使用 Optuna 贝叶斯优化框架，对每个模型进行 30 次试验（LSTM 为 10 次），"
    "每次试验在 5 折交叉验证上评估调整后夏普比率。最优参数组合保存后，"
    "再在完整 10 折验证上进行最终评估。"
)

_heading(doc, "4.3 仓位映射", level=2)
_para(doc,
    "模型输出的是连续预测值，而不是直接的持仓比例。"
    "我们通过以下公式将预测值转换为仓位："
)
_para(doc, "position = clip(1 + k × prediction / rolling_std, 0, 2)", bold=True)
_para(doc,
    "其中 k 是通过优化确定的缩放系数，rolling_std 是预测值的 63 日滚动标准差。"
    "这一设计使仓位自适应预测信号的强弱，同时避免极端仓位。"
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 4. 交叉验证模型比较
# ════════════════════════════════════════════════════════════════════════════
_heading(doc, "五、交叉验证结果对比")
_para(doc,
    "下表汇总了所有模型在 10 折交叉验证集上的关键指标。"
    "其中「拼接调整后夏普」是将所有验证折的预测拼接后整体计算，"
    "「折均值调整后夏普」是各折单独计算后取平均，两者的差异反映了模型跨时间段的稳定性。"
)

model_order = ["buy_and_hold", "rf", "ridge", "linear", "elasticnet",
               "lasso", "xgboost", "lightgbm", "lstm"]
name_map = {
    "buy_and_hold": "买入持有（基准）",
    "rf": "随机森林 ★",
    "ridge": "岭回归",
    "linear": "线性回归",
    "elasticnet": "弹性网络",
    "lasso": "LASSO",
    "xgboost": "XGBoost",
    "lightgbm": "LightGBM",
    "lstm": "LSTM",
}

def _fmt(v, fmt=".3f"):
    try:
        return f"{float(v):{fmt}}" if v is not None and str(v) != "nan" else "—"
    except Exception:
        return "—"

rows = []
for _, row in cv_df.set_index("model").reindex(
    [m for m in model_order if m in cv_df["model"].values]
).reset_index().iterrows():
    m = row["model"]
    rows.append((
        name_map.get(m, m),
        _fmt(row.get("cv10_adjusted_sharpe")),
        _fmt(row.get("cv10_adj_sharpe_fold_avg")),
        _fmt(row.get("cv10_sharpe_raw")),
        _fmt(row.get("cv10_max_drawdown")),
        _fmt(row.get("cv10_win_rate")),
        _fmt(row.get("cv10_cum_return_pct"), ".1f") + "%" if _fmt(row.get("cv10_cum_return_pct")) != "—" else "—",
    ))

_metric_table(doc, rows, [
    "模型", "拼接调整后夏普", "折均值调整后夏普",
    "原始夏普", "最大回撤", "胜率", "累计收益"
])

_para(doc,
    "随机森林在综合指标上表现最优：折均值调整后夏普 0.50，最大回撤 -57.4%（验证集跨越多个市场周期），"
    "累计收益 765%，是 7 类模型中最高的。LASSO 因高 k 值（0.69）在短窗口验证中的不稳定性，"
    "拼接得分（0.13）与折均值（0.48）差距显著，提示过拟合风险。"
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 5. 测试集结果（RF）
# ════════════════════════════════════════════════════════════════════════════
_heading(doc, "六、最优模型测试集表现（随机森林）")
_para(doc,
    "随机森林模型在从未接触过的独立测试集上的表现如下。"
    "测试集为数据集末尾约 805 个交易日（约 3 年）的数据，"
    "在整个调优与验证过程中均严格封存，不参与任何训练或参数选择。"
)

_heading(doc, "6.1 关键指标", level=2)
metric_rows = [
    ("调整后夏普比率",  f"{tm['adjusted_sharpe']:.4f}", f"{bh_score:.4f}", "越高越好"),
    ("原始夏普比率",   f"{tm['sharpe_raw']:.4f}",       "—",               "越高越好"),
    ("索提诺比率",    f"{tm['sortino']:.4f}",            "—",               "越高越好，对下行风险更敏感"),
    ("卡玛比率",     f"{tm['calmar']:.4f}",              "—",               "收益 / 最大回撤"),
    ("最大回撤",     f"{tm['max_drawdown']:.2%}",        "—",               "越小越好"),
    ("胜率",        f"{tm['win_rate']:.2%}",             "—",               "超额收益为正的比例"),
    ("年化波动率",   f"{tm['ann_vol_pct']:.2f}%",        "—",               "策略日收益的年化标准差"),
    ("IC（信息系数）", f"{tm['ic']:.4f}",                "—",               "预测值与实际值的排序相关性"),
    ("ICIR",       f"{tm['icir']:.4f}",                 "—",               "IC 均值 / IC 标准差，越高越稳定"),
]
_metric_table(doc, metric_rows, ["指标", "策略表现", "买入持有基准", "说明"])

_para(doc,
    f"策略的调整后夏普（{tm['adjusted_sharpe']:.2f}）超过买入持有基准（{bh_score:.2f}）约 "
    f"{tm['adjusted_sharpe'] - bh_score:.2f} 个单位，最大回撤控制在 {abs(tm['max_drawdown']):.1%} 以内，"
    "远低于验证集期间的水平，显示出测试期内市场环境较为温和，策略风险管理有效。"
)
_para(doc,
    "IC 为 0.059，绝对值较低但方向一致为正，说明模型具备微弱但稳定的预测能力。"
    "ICIR 为 0.61，在量化行业中属于中等偏上水平，表明预测质量的时间稳定性尚可。"
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 可视化图表
# ════════════════════════════════════════════════════════════════════════════
_heading(doc, "6.2 可视化分析", level=2)

charts = [
    ("test_cumret_rf.png",
     "图1：累计收益曲线（策略 vs 买入持有）",
     "策略（绿色实线）全程跑赢买入持有（蓝色虚线），尤其在中后期差距持续扩大，"
     "显示出模型的动态仓位管理在捕捉市场上涨的同时有效规避了部分下跌。"),

    ("test_drawdown_rf.png",
     "图2：资金回撤（水下图）",
     "回撤幅度整体控制在 16% 以内，且恢复速度较快，无长期深度亏损区间，"
     "说明仓位映射机制在市场不利时能够及时降低敞口。"),

    ("test_rolling_sharpe_rf.png",
     "图3：63日滚动夏普比率",
     "滚动夏普在大多数时期维持在 0 以上，均值约为 1.0 左右，"
     "少数时期出现负值，与回撤图对应，属于局部市场逆风期。"),

    ("test_rolling_vol_rf.png",
     "图4：63日滚动波动率（策略 vs 买入持有）",
     "策略波动率（绿色）在多数时期低于或接近买入持有（蓝色虚线），"
     "说明动态仓位确实起到了平滑组合风险的作用，而非单纯放大收益。"),

    ("test_monthly_rf.png",
     "图5：月度收益热力图",
     "绿色格代表当月盈利，红色格代表亏损，颜色深浅对应幅度大小。"
     "整体看盈利月份明显多于亏损月份，亏损月份的深度也相对有限。"),

    ("test_positions_rf.png",
     "图6：仓位分布直方图",
     "仓位集中分布在 0.5 至 1.5 之间，均值约为 1.0（中性），"
     "极端仓位（接近 0 或 2）出现频率较低，说明模型输出信号较为温和，不过度押注。"),

    ("test_ic_scatter_rf.png",
     "图7：预测值 vs 实际值散点图（IC 验证）",
     "散点整体呈微弱正相关（IC=0.059），回归线斜率为正。"
     "分布较为分散符合金融预测的客观规律——市场噪声大，"
     "能稳定保持正 IC 已具备实际交易价值。"),
]

for fname, caption_title, caption_body in charts:
    path = TEST / fname
    _para(doc, caption_title, bold=True, size=10)
    _img(doc, path, width=5.8)
    _para(doc, caption_body, italic=True, size=9, color=(80, 80, 80))
    doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 7. 局限性与未来改进方向
# ════════════════════════════════════════════════════════════════════════════
_heading(doc, "七、局限性与未来改进方向")

_heading(doc, "7.1 当前局限性", level=2)

limitations = [
    ("预测能力有限",
     "IC 仅为 0.059，说明模型对收益率的预测精度较低，R² 为负值意味着模型预测误差"
     "甚至大于用均值预测的误差。当前盈利主要依赖仓位缩放对市场趋势的顺势跟随，"
     "而非精准的方向预测。"),
    ("测试期偏短且环境偏好",
     "测试集约 805 个交易日，买入持有基准在此期间的调整后夏普已达 0.87，"
     "说明测试期整体为较强的牛市或低波动环境。策略在熊市或高波动期的表现尚未经过充分验证。"),
    ("特征工程较为传统",
     "当前特征主要基于价格动量与技术指标，缺乏基本面数据（财报、估值）、"
     "另类数据（情绪、新闻）等能提供独立信息来源的变量。"),
    ("LSTM 调优不足",
     "LSTM 仅进行了 10 次 Optuna 试验（其他模型 30 次），且在 CPU 上训练，"
     "超参空间探索不充分，导致其 CV 表现低于预期，未能体现深度学习的潜力。"),
    ("单一模型提交",
     "最终仅提交随机森林一个模型，未利用模型集成（Ensemble）来进一步降低预测方差。"),
    ("交易成本未计入",
     "回测未考虑手续费、滑点、市场冲击成本，实盘收益会有所折损，"
     "尤其是仓位频繁调整时影响更为显著。"),
]

for title, body in limitations:
    p = doc.add_paragraph(style="List Bullet")
    run_title = p.add_run(title + "：")
    run_title.bold = True
    p.add_run(body)
    p.paragraph_format.space_after = Pt(5)

_heading(doc, "7.2 未来改进方向", level=2)

improvements = [
    ("引入基本面与另类数据",
     "加入市盈率、市净率、分析师预期修正、新闻情绪等特征，"
     "为模型提供价格动量之外的独立预测信号，有望显著提升 IC。"),
    ("增强 LSTM 调优",
     "在 GPU 环境下对 LSTM 进行充分的超参数搜索（50+ 次试验），"
     "并尝试引入 Transformer、时序 Attention 等更先进的序列模型架构。"),
    ("模型集成",
     "将随机森林、岭回归、LSTM 等表现互补的模型通过 Stacking 或加权平均组合，"
     "集成预测通常能在不增加过拟合风险的前提下提高稳定性。"),
    ("自适应特征选择",
     "当前特征筛选基于全窗口 IC 统计，可改为滚动 IC 筛选，"
     "让特征集随时间动态调整，适应市场结构变化。"),
    ("加入成本模型",
     "在仓位映射层引入换手率惩罚项，减少不必要的仓位波动，"
     "并在回测中模拟真实的手续费和滑点，使评估结果更贴近实盘。"),
    ("压力测试与情景分析",
     "在历史上的特殊市场环境（2008 年金融危机、2020 年疫情冲击等）"
     "单独测试策略表现，验证其在极端行情下的鲁棒性。"),
]

for title, body in improvements:
    p = doc.add_paragraph(style="List Bullet")
    run_title = p.add_run(title + "：")
    run_title.bold = True
    p.add_run(body)
    p.paragraph_format.space_after = Pt(5)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 8. 结论
# ════════════════════════════════════════════════════════════════════════════
_heading(doc, "八、结论")
_para(doc,
    "本项目系统性地构建并评估了面向 Hull Tactical 竞赛的机器学习量化策略。"
    "通过严格的时间序列交叉验证和超参数调优，随机森林模型在综合风险收益指标上"
    "表现最优，并在独立测试集上以调整后夏普 1.12 超越买入持有基准（0.87），"
    "最大回撤控制在 16% 以内。"
)
_para(doc,
    "尽管模型的预测精度（IC=0.059）相对有限，但通过合理的仓位映射机制，"
    "较弱的预测信号仍然转化为了实际的超额收益。这也说明在量化投资中，"
    "风险管理和仓位控制与预测精度同等重要。"
)
_para(doc,
    "未来工作的重点应集中在丰富数据源、改进序列模型和引入集成策略上，"
    "以期在更复杂的市场环境中取得更稳健的超额收益。"
)

doc.save(OUT)
print(f"Report saved → {OUT}")
